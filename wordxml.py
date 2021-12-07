import os
import zipfile
import pandas  # +openpyxl
import docx
from docx.oxml.shared import qn  # feel free to move these out
from docx.oxml.xmlchemy import OxmlElement
from tkinter import filedialog
import sys

print("ворд")
pathword = filedialog.askopenfilename()
pathwork = os.path.dirname(pathword)
pathzip = pathwork + "/B.zip"
pathword2 = pathwork + "/Шаблон написания справки.xlsx"
mem=0
isch=0
usch=0
found=""
df = pandas.read_excel(pathword2)


def resource_path(relative):
    if hasattr(sys, "_MEIPASS"):
        return os.path.join(sys._MEIPASS, relative)
    return os.path.join(relative)


def _set_cell_background(cell, fill, color=None, val=None):
    """
    @fill: Specifies the color to be used for the background
    @color: Specifies the color to be used for any foreground
    pattern specified with the val attribute
    @val: Specifies the pattern to be used to lay the pattern
    color over the background color.
    """
    cell_properties = cell._element.tcPr
    try:
        cell_shading = cell_properties.xpath('w:shd')[0]  # in case there's already shading
    except IndexError:
        cell_shading = OxmlElement('w:shd') # add new w:shd element to it
    if fill:
        cell_shading.set(qn('w:fill'), fill)  # set fill property, respecting namespace
    if color:
        pass # TODO
    if val:
        pass # TODO
    cell_properties.append(cell_shading)  # finally extend cell props with shading element

#w:fill="FFFFFF" w:themeFill="background1"

def info_update(doc, old_info, new_info, colorr):  # Paint cells
    outtt = 0
    if colorr == "":
        colorr = "FFFFFF"
    for table in doc.tables:
        r = 0
        for row in table.rows:
            c = 0
            for cell in row.cells:
                prov=cell.text
                if old_info == prov:  # cell.text == old_info:
                    outtt = 1
                    _set_cell_background(table.rows[r].cells[c], colorr)  # есть куда ускорять...........................
                    #cell._tc.get_or_add_tcPr().append(shading_elm_1) # информация о замене
                if outtt == 1:
                    return
                c=c+1
            r=r+1


def info_update1(doc, old_info, new_info):  # Delete stroke
    outt=0
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text == old_info:
                outt = 1
                if new_info == "":
                    delete_paragraph(para)
            run.text = run.text.replace(old_info, new_info)  # информация о замене
            if outt == 1:
                return


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

doc = docx.Document(pathword)
for i in range(50, len(df.index)): #len(df.index)
    if not (str(df.iloc[i, 2]) == "nan" or str(df.iloc[i, 2]) == "FFFFFF"):
        info_update(doc, str(df.iloc[i, 0]), str(df.iloc[i, 1]), str(df.iloc[i, 2]))
        #print(df.iloc[i, 0])
        print(i, "--", df.iloc[i, 2])
for i in range(1, 54): #len(df.index)
    if not str(df.iloc[i, 1]) == "nan":
        info_update1(doc, str(df.iloc[i, 0]), str(df.iloc[i, 1]))
    else:
        info_update1(doc, str(df.iloc[i, 0]), "")
    print(i)

doc.save(pathwork + "/B.docx")
os.rename(pathwork + "/B.docx", pathwork + "/B.zip")

fantasy_zip = zipfile.ZipFile(pathzip)  # extract zip (+need rename docx to zip +need raname vise versa
fantasy_zip.extractall(pathwork + "/B")
fantasy_zip.close()

with open(pathwork + "/B/word/document.xml", 'r', encoding='utf-8') as f:  # save before chenge
    get_all = f.readlines()

print("xml opened")
with open(pathwork + "/B/word/document.xml", 'w', encoding='utf-8') as f:  # look for { and chenge it
    for i in get_all:         # STARTS THE NUMBERING FROM 1 (by default it begins with 0)
        usch=len(i)-1
        for u in i:
            if get_all[isch][usch] == "}":
                mem = 1
            if mem == 1:
                found = get_all[isch][usch]+found
            if get_all[isch][usch] == "{":
                mem = 0
                for dfn in range(0, len(df.index)):  # look for found in df
                    if str(df.iloc[dfn, 0]) == found:
                        if str(df.iloc[dfn, 1]) == "nan":
                            df.iloc[dfn, 1] = ""
                        get_all[isch] = get_all[isch][:usch] + str(df.iloc[dfn, 1]) + get_all[isch][usch+len(found):]
                found = ""
            usch = usch - 1
        isch = isch + 1
    f.writelines(get_all)  # save it
print("XML chanched")
fantasy_zip = zipfile.ZipFile(pathwork + "/B.zip", 'w')
for folder, subfolders, files in os.walk(pathwork + "/B"):
    for file in files:
        fantasy_zip.write(os.path.join(folder, file), os.path.relpath(os.path.join(folder, file), pathwork + "/B"))
fantasy_zip.close()  # transform it to zip
print("zip saved")
try:
    os.remove(pathwork + "/" + str(df.iloc[696, 1]) + ".docx")
except:
    asd = 1
os.rename(pathwork + "/B.zip", pathwork + "/" + str(df.iloc[696, 1]) + ".docx")




# mf2 = open("items2.xml", "w")
# mf2.write(nf)
#myzip.write("word/document.xml", "C:/Users/IMatveev/Desktop/ппворд/A.zip\\test.py", zipfile.ZIP_DEFLATED )


